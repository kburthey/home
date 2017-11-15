#! Python3

from jinja2 import Environment, FileSystemLoader
import docx
import os, sys, openpyxl, pprint, random
from openpyxl.styles import Font, PatternFill, Border, Side
from time import strftime
from openpyxl.chart import (
    Reference,
    Series,
    BarChart3D,
    BarChart
    )

''' Column Headers from EZPlan Output file
PI ID & Title: Col 4,
Activity ID: Col 5,
PI ID: Col 6,
TK ID: Col 7,
Work Type: Col 8,
Crop: Col 13,
Planned SD: Col 16,
Allocated SD: Col 17,
Planned $: Col 20,
Allocated $: Col 21,
Named Resource: Col 22,
Activity Comments: Col 23,
Portfolio: Col 24
'''

def color():
    hecks = "ABCDEF0123456789"
    randomColor = ""
    for i in range(0, 6):
        randomColor += (hecks[random.randint(0, len(hecks)-1)])
    return randomColor
rColor = color()

PATH = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_ENVIRONMENT = Environment(
    autoescape=False,
    loader=FileSystemLoader(os.path.join(PATH, 'templates')),
    trim_blocks=False)

def render_template(template_file, context):
    return TEMPLATE_ENVIRONMENT.get_template('index.html').render(context)

doc = docx.Document('demo.docx')
sentence = []
for i in range(0, len(doc.paragraphs)-1):
    sentence.append(doc.paragraphs[i].text)

def create_index_html():
    fname = "output.html"
    title = sys.argv[1]
    urls = ['http://example.com/1', 'http://example.com/2', 'http://example.com/3']
    text = ['hello', 'world']
    apples = {
        'red': 'sweet',
        'green': 'bitter'
        }
    context = {
        'links':urls,
        'text':text,
        'dict': apples,
        'sentence': sentence,
        'title': title,
        'color': rColor
        }
    #
    with open(fname, 'w') as f:
        html = render_template('index.html', context)
        f.write(html)


def main():
    create_index_html()

######

if __name__ == "__main__":
    main()

    
