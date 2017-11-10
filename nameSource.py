import openpyxl
import docx
import sys
import os
'''
PI ID & Title: Col 4,
Activity ID: Col 5,
PI ID: Col 6,
TK ID: Col 7,
Work Type: Col 8,
Crop: Col 13,
Planned SD: Col 16,
Allocated SD: Col 17,
Planned $: Col 20,
Allocated $: Col 21,
Named Resource: Col 22,
Activity Comments: Col 23,
Portfolio: Col 24
'''

sourceFile = sys.argv[1]

nameSource = openpyxl.load_workbook(sourceFile)
sheets = nameSource.get_sheet_names()
source = sheets[0]
page = nameSource.get_sheet_by_name(source)
rowMax = page.max_row


kirkreport = docx.Document()
kirkTotal = []
kirkreport.add_heading('Kirkland Burthey: 2017 Activities', 0)
for i in range(1,rowMax):
    if str(page.cell(row=i, column=22).value) == 'Burthey, Kirkland': #and str(page.cell(row=i, column=8).value) == 'Seq BSA':
        kirkreport.add_paragraph("Activity ID: " + str(page.cell(row=i, column=5).value) + " \nSD: " + str(page.cell(row=i, column=17).value) + " \nWork Type: " + str(page.cell(row=i, column=8).value) + " \nActivity Comment: " + str(page.cell(row=i, column=23).value) + "\n\n")
        kirkTotal.append(page.cell(row=i, column=17).value)
kirkgrandTotal = sum(kirkTotal)
kirkreport.add_paragraph('Total allocated SD: ' + str(kirkgrandTotal))      
kirkreport.save('c:\\Users\\t862537\\kirknamedSource.docx')

laurareport = docx.Document()
lauraTotal = []
laurareport.add_heading('Laura Kavanaugh: 2017 Activities', 0)
for i in range(1,rowMax):
    if str(page.cell(row=i, column=22).value) == 'Kavanaugh, Laura':
        laurareport.add_paragraph("Activity ID: " + str(page.cell(row=i, column=5).value) + " \nSD: " + str(page.cell(row=i, column=17).value) + " \nWork Type: " + str(page.cell(row=i, column=8).value) + " \nActivity Comment: " + str(page.cell(row=i, column=23).value) + "\n\n")
        lauraTotal.append(page.cell(row=i, column=17).value)
lauragrandTotal = sum(lauraTotal)
laurareport.add_paragraph('Total allocated SD: ' + str(lauragrandTotal))      
laurareport.save('c:\\Users\\t862537\\lauranamedSource.docx')


joshreport = docx.Document()
joshTotal = []
joshreport.add_heading('Josh Cohn: 2017 Activities', 0)
for i in range(1,rowMax):
    if str(page.cell(row=i, column=22).value) == 'Cohn, Jonathan':
        joshreport.add_paragraph("Activity ID: " + str(page.cell(row=i, column=5).value) + " \nSD: " + str(page.cell(row=i, column=17).value) + " \nWork Type: " + str(page.cell(row=i, column=8).value) + " \nActivity Comment: " + str(page.cell(row=i, column=23).value) + "\n\n")
        joshTotal.append(page.cell(row=i, column=17).value)
joshgrandTotal = sum(joshTotal)
joshreport.add_paragraph('Total allocated SD: ' + str(joshgrandTotal))      
joshreport.save('c:\\Users\\t862537\\joshnamedSource.docx')
