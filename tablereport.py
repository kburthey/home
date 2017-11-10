import openpyxl
import docx
import sys
import os
'''
PI ID & Title: Col 4,
Activity ID: Col 5,
PI ID: Col 6,
TK ID: Col 7,
Work Type: Col 8,
Crop: Col 13,
Planned SD: Col 16,
Allocated SD: Col 17,
Planned $: Col 20,
Allocated $: Col 21,
Named Resource: Col 22,
Activity Comments: Col 23,
Portfolio: Col 24
'''

sourceFile = sys.argv[1]

nameSource = openpyxl.load_workbook(sourceFile)
sheets = nameSource.get_sheet_names()
source = sheets[0]
page = nameSource.get_sheet_by_name(source)
rowMax = page.max_row


kirkreport = docx.Document()
kirkTotal = []
kirkreport.add_heading('Kirkland Burthey: 2017 Activities', 0)
for i in range(1,rowMax):
    if str(page.cell(row=i, column=22).value) == 'Burthey, Kirkland': #and str(page.cell(row=i, column=8).value) == 'Seq BSA':
        kirkreport.add_paragraph("Activity ID: " + str(page.cell(row=i, column=5).value) + " \nSD: " + str(page.cell(row=i, column=17).value) + " \nWork Type: " + str(page.cell(row=i, column=8).value) + " \nActivity Comment: " + str(page.cell(row=i, column=23).value) + "\n\n")
        kirkTotal.append(page.cell(row=i, column=17).value)
kirkgrandTotal = sum(kirkTotal)
kirkreport.add_paragraph('Total allocated SD: ' + str(kirkgrandTotal))      
kirkreport.save('c:\\Users\\t862537\\kirknamedSource.docx')

kirktable = docx.Document()
table = kirktable.add_table(rows=1, cols=4)
#table.style = 'TableGrid'
heading_cells = table.rows[0].cells
heading_cells[0].text = str(page.cell(row=1, column=5).value)
heading_cells[1].text = str(page.cell(row=1, column=17).value)
heading_cells[2].text = str(page.cell(row=1, column=8).value)
heading_cells[3].text = str(page.cell(row=1, column=23).value)
for i in range(2, rowMax):
    if str(page.cell(row=i, column=22).value) == 'Burthey, Kirkland':
        cells = table.add_row().cells    
        cells[0].text = str(page.cell(row=i, column=5).value)
        cells[1].text = str(page.cell(row=i, column=17).value)
        cells[2].text = str(page.cell(row=i, column=8).value)
        cells[3].text = str(page.cell(row=i, column=23).value)
kirktable.save('kirktable.docx')
